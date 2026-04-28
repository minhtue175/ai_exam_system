import random
import copy
import io
from docx import Document
from docx.shared import Pt, RGBColor
from django.template.loader import render_to_string
from weasyprint import HTML

def prepare_quiz_data(quiz, should_shuffle=False):
    """
    Hàm chuẩn bị dữ liệu đề thi siêu tốc độ.
    - Ép Database trả về Dict (không dùng Object).
    - Fix triệt để bug xáo trộn đáp án trùng lặp.
    """

    questions_raw = list(quiz.questions.values(
        'id', 'question_text', 'options', 'correct_answer', 'explanation'
    ))
    
    rng = random.Random()

    if should_shuffle:
        rng.shuffle(questions_raw)

    prepared_data = []
    
    
    for q in questions_raw:
        options = list(q['options']) 
        correct_index = q['correct_answer']
        
        if should_shuffle:
            indices = list(range(len(options))) 
            rng.shuffle(indices) 
            
            options = [options[i] for i in indices]
            correct_index = indices.index(correct_index)
            
        prepared_data.append({
            'id': q['id'],
            'question_text': q['question_text'],
            'options': options,
            'correct_index': correct_index,
            'explanation': q['explanation'] or 'Không có giải thích',
        })

    return prepared_data

def generate_word_document(quiz_title, prepared_data, mode='student'):
    """
    Hàm tạo file Word từ dữ liệu đã được chuẩn bị.
    mode: 'student' (chỉ có đề) hoặc 'teacher' (có đáp án đỏ + giải thích)
    """
    doc = Document()
    

    title = doc.add_heading(f'ĐỀ THI: {quiz_title.upper()}', 0)
    title.alignment = 1 
    
    if mode == 'teacher':
        subtitle = doc.add_paragraph('BẢN DÀNH CHO GIÁO VIÊN (CÓ ĐÁP ÁN CHI TIẾT)')
        subtitle.alignment = 1
        
    doc.add_paragraph()
    

    labels = ['A', 'B', 'C', 'D']
    
    for i, q in enumerate(prepared_data, 1):
        
        p_question = doc.add_paragraph()
        run_q = p_question.add_run(f'Câu {i}: {q["question_text"]}')
        run_q.bold = True
        
        # In 4 đáp án
        for j, option_text in enumerate(q['options']):
            p_option = doc.add_paragraph()
            run_opt = p_option.add_run(f'   {labels[j]}. {option_text}')
            
            
            if mode == 'teacher' and j == q['correct_index']:
                run_opt.bold = True
                run_opt.font.color.rgb = RGBColor(255, 0, 0) # Màu đỏ
                
        
        if mode == 'teacher' and q['explanation']:
            p_exp = doc.add_paragraph()
            run_exp = p_exp.add_run(f'   Giải thích: {q["explanation"]}')
            run_exp.italic = True
            run_exp.font.color.rgb = RGBColor(89, 89, 89) 
            
        doc.add_paragraph() 
        
   
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def generate_pdf_document(quiz_title, prepared_data, mode='student'):
    """
    Hàm tạo file PDF bằng cách render file HTML rồi convert sang PDF.
    """
    
    context = {
        'title': quiz_title,
        'questions': prepared_data,
        'mode': mode
    }
    
   
    html_string = render_to_string('quizzes/export_pdf_template.html', context)
    
  
    buffer = io.BytesIO()
    HTML(string=html_string).write_pdf(buffer)
    buffer.seek(0)
    
    return buffer